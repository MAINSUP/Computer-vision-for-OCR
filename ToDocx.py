from docx import Document
from docx.enum.section import WD_SECTION
# #*#import easyocr
# #*#import torch
import pytesseract
import TextStyle
from docx.shared import Pt
from stqdm import stqdm
import os
import warnings
warnings.filterwarnings('ignore')
import locale
locale.getpreferredencoding()


for r, s, f in os.walk("/"):
    for i in f:
        if "tesseract" in i:
            os.path.join(r, i)
pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'
# tessdata_dir_config = "--tessdata-dir 'C:\\Program Files \\Tesseract-OCR\\tessdata\\"


def im2dox(file, language, confidence):
    document = Document()  # creating Word document instance
    # #*# torch.cuda.empty_cache()
    # #*# reader = easyocr.Reader([language], gpu=False)  # calling ocr reader

    print("Extracting text to docx file")
    for ref, page in enumerate(file):
        print("Reading page {}".format(ref+1))
        for segment in stqdm(page):
            p = document.add_paragraph()  # for each text layout segment, a separate paragraph is created
            # #*#result = reader.readtext(segment)  # and filled with predicted text
            result = pytesseract.image_to_string(segment, lang=language, config='--psm 6')
            font_style = TextStyle.text_style(segment, confidence)
            # #*#for (bbox, text, prob) in result:
            for t in result:
                try:
                    text = t.encode("cp850").decode("cp850")
                except UnicodeEncodeError:
                    print(t)
                if font_style != 'Cursive':
                    try:
                        run = p.add_run(text)
                        run.font.size = Pt(14)
                    except ValueError:
                        del text
                else:
                     try:
                         run = p.add_run(text)
                         run.font.name = 'Brush Script MT'       # writing cursive font text
                         run.font.size = Pt(14)
                     except ValueError:
                        del text
    document.add_section(start_type=WD_SECTION.NEW_PAGE)
    return document


def txt2dox(text):
    document = Document()  # creating Word document instance
    p = document.add_paragraph()
    p.add_run(text)    # writing standard font text
    return document

# #*# for use with easyocr lib
